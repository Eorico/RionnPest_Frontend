from PyQt5.QtWidgets import (
    QMainWindow, QListWidgetItem, QMessageBox, QPushButton,
    QHBoxLayout, QWidget, QTableWidgetItem, QFileDialog, QMenu, QAction
)
from PyQt5.QtCore import Qt, QSize
from PyQt5.QtGui import QCursor, QFont
from ui.docs_viewer import Ui_DocxViewer

# ── Zoom constants ────────────────────────────────────────────────────────────
_ZOOM_MIN  = 50
_ZOOM_MAX  = 200
_ZOOM_STEP = 10

# ── Template validation keywords (at least 3 must match in header row) ────────
_TEMPLATE_KEYWORDS = {"date", "client", "time", "chemical", "category", "remarks", "actual"}
_TEMPLATE_MIN_COLS = 6


class DocxViewer(QMainWindow):

    def __init__(self, api_service, parent=None):
        super().__init__(parent)
        self.api = api_service
        self.ui  = Ui_DocxViewer()
        self.ui.setupUi(self)

        self._doc_map     = {}
        self._current_doc = None
        self._zoom        = 100          # current zoom %
        self._raw_bytes   = None         # bytes of currently loaded doc

        # Correct table headers (7 columns)
        self.ui.inventoryTable1.setColumnCount(7)
        for col, txt in enumerate([
            "Category", "Date", "Name of Client", "Time",
            "Chemical/s Used", "Actual Chemical/s Used", "Remarks"
        ]):
            self.ui.inventoryTable1.setHorizontalHeaderItem(col, QTableWidgetItem(txt))

        self._connect_signals()
        self.refresh_file_list()

    # ── Signal wiring ─────────────────────────────────────────────────────────

    def _connect_signals(self):
        self.ui.filesList.itemClicked.connect(self._on_file_selected)
        self.ui.tocList.itemClicked.connect(self._on_toc_selected)

        # Zoom
        self.ui.actionZoomIn.triggered.connect(self._zoom_in)
        self.ui.actionZoomOut.triggered.connect(self._zoom_out)
        self.ui.actionZoomReset.triggered.connect(self._zoom_reset)

        # Ctrl+scroll wheel zoom
        self.ui.pageScrollArea.installEventFilter(self)

        # Open / Export / Print
        self.ui.actionOpen.triggered.connect(self._open_file)
        self.ui.actionExportPDF.triggered.connect(self._export_pdf)
        self.ui.actionExportDocx.triggered.connect(self._export_docx)
        self.ui.actionPrint.triggered.connect(self._print_document)
        self.ui.actionClose.triggered.connect(self.close)
        self.ui.actionRefresh.triggered.connect(self.refresh_file_list)

        # Export button shows a small menu
        self.ui.btnExport.clicked.disconnect()          # detach from actionExportPDF
        self.ui.btnExport.clicked.connect(self._show_export_menu)

    def eventFilter(self, obj, event):
        from PyQt5.QtCore import QEvent
        if obj is self.ui.pageScrollArea and event.type() == QEvent.Wheel:
            if event.modifiers() & Qt.ControlModifier:
                if event.angleDelta().y() > 0:
                    self._zoom_in()
                else:
                    self._zoom_out()
                return True
        return super().eventFilter(obj, event)

    # ── Zoom ─────────────────────────────────────────────────────────────────

    def _zoom_in(self):
        if self._zoom < _ZOOM_MAX:
            self._zoom = min(self._zoom + _ZOOM_STEP, _ZOOM_MAX)
            self._apply_zoom()

    def _zoom_out(self):
        if self._zoom > _ZOOM_MIN:
            self._zoom = max(self._zoom - _ZOOM_STEP, _ZOOM_MIN)
            self._apply_zoom()

    def _zoom_reset(self):
        self._zoom = 100
        self._apply_zoom()

    def _apply_zoom(self):
        factor = self._zoom / 100.0

        # Update zoom label
        self.ui.zoomLabel.setText(f"{self._zoom}%")

        # Scale page widths
        base_min_w  = 900
        base_max_w  = 1200
        self.ui.paperPage1.setMinimumWidth(int(base_min_w * factor))
        self.ui.paperPage2.setMinimumWidth(int(base_min_w * factor))
        self.ui.pageInnerContainer.setMaximumWidth(int(base_max_w * factor))

        # Scale table content font
        base_pt = 10
        scaled_pt = max(7, int(base_pt * factor))
        tbl_font = QFont("Segoe UI", scaled_pt)
        self.ui.inventoryTable1.setFont(tbl_font)

        # Scale row heights
        base_row_h = 40
        self.ui.inventoryTable1.verticalHeader().setDefaultSectionSize(
            max(28, int(base_row_h * factor)))

        # Scale table header font
        hdr_font = QFont("Segoe UI", max(7, int(9 * factor)), QFont.DemiBold)
        self.ui.inventoryTable1.horizontalHeader().setFont(hdr_font)

        # Re-fit rows to new font
        self.ui.inventoryTable1.resizeRowsToContents()

    # ── Open file (with template validation) ─────────────────────────────────

    def _open_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Open Document", "",
            "Word Documents (*.docx);;All Files (*)"
        )
        if not path:
            return

        try:
            with open(path, 'rb') as f:
                raw = f.read()
        except Exception as e:
            QMessageBox.critical(self, "Open Failed", f"Could not read file:\n{e}")
            return

        valid, reason = self._validate_template(raw)
        if not valid:
            QMessageBox.warning(
                self, "Invalid Document Template",
                f"This file does not match the Raionn template.\n\n"
                f"Reason: {reason}\n\n"
                f"Expected: a table with columns for Category, Date, Client, "
                f"Time, Chemicals Used, Actual Chemicals Used, and Remarks."
            )
            return

        import os
        file_name = os.path.basename(path)
        records = self._parse_docx_to_records(raw)
        self._raw_bytes = raw
        self._current_doc = {"id": None, "title": file_name, "records": records}
        self._render_document(records, file_name)
        self._build_toc(records)
        self.ui.docTitleLabel.setText(file_name)

    @staticmethod
    def _validate_template(raw: bytes) -> tuple[bool, str]:
        """
        Returns (True, "") if the docx matches the Raionn template,
        or (False, reason_string) if it does not.
        """
        from io import BytesIO
        from docx import Document
        try:
            doc = Document(BytesIO(raw))
        except Exception as e:
            return False, f"Could not parse file: {e}"

        if not doc.tables:
            return False, "No tables found in the document"

        tbl = doc.tables[0]
        if len(tbl.columns) < _TEMPLATE_MIN_COLS:
            return False, (f"Table has only {len(tbl.columns)} column(s); "
                           f"expected at least {_TEMPLATE_MIN_COLS}")

        if not tbl.rows:
            return False, "Table has no rows"

        header_text = " ".join(c.text.strip().lower() for c in tbl.rows[0].cells)
        matched = sum(1 for kw in _TEMPLATE_KEYWORDS if kw in header_text)
        if matched < 3:
            return False, (f"Header row does not look like the Raionn template "
                           f"(matched {matched}/{len(_TEMPLATE_KEYWORDS)} expected keywords)")

        return True, ""

    # ── Export ────────────────────────────────────────────────────────────────

    def _show_export_menu(self):
        menu = QMenu(self)
        menu.setStyleSheet("""
            QMenu { background: #fff; border: 1px solid #E0EDE6;
                    border-radius: 8px; font: 10pt 'Segoe UI'; padding: 4px 0; }
            QMenu::item { padding: 7px 22px 7px 16px; border-radius: 4px; margin: 1px 4px; }
            QMenu::item:selected { background: #2D6A4F; color: #fff; }
        """)
        pdf_action  = menu.addAction("📄  Export as PDF")
        docx_action = menu.addAction("📝  Save as DOCX")
        action = menu.exec_(self.ui.btnExport.mapToGlobal(
            self.ui.btnExport.rect().bottomLeft()))
        if action == pdf_action:
            self._export_pdf()
        elif action == docx_action:
            self._export_docx()

    def _export_pdf(self):
        if not self._current_doc or not self._current_doc.get("records"):
            QMessageBox.information(self, "No Document", "Load a document first.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "Export as PDF", 
            self._current_doc.get("title", "report").replace(".docx", "") + ".pdf",
            "PDF Files (*.pdf)"
        )
        if not path:
            return

        try:
            from PyQt5.QtPrintSupport import QPrinter
            from PyQt5.QtGui import QTextDocument

            printer = QPrinter(QPrinter.HighResolution)
            printer.setOutputFormat(QPrinter.PdfFormat)
            printer.setOutputFileName(path)
            printer.setPageSize(QPrinter.A4)
            printer.setPageMargins(20, 20, 20, 20, QPrinter.Millimeter)

            html = self._build_html_for_print()
            doc = QTextDocument()
            doc.setHtml(html)
            doc.print_(printer)

            QMessageBox.information(self, "Exported", f"PDF saved to:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Failed", f"Could not export PDF:\n{e}")

    def _export_docx(self):
        if not self._current_doc:
            QMessageBox.information(self, "No Document", "Load a document first.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "Save as DOCX",
            self._current_doc.get("title", "report"),
            "Word Documents (*.docx)"
        )
        if not path:
            return

        try:
            # If opened from API, re-download; if opened locally, use raw bytes
            if self._current_doc.get("id") is not None:
                raw = self.api.download_document(self._current_doc["id"])
            else:
                raw = self._raw_bytes

            if not raw:
                QMessageBox.critical(self, "Error", "Document data not available.")
                return

            with open(path, 'wb') as f:
                f.write(raw)
            QMessageBox.information(self, "Saved", f"Document saved to:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Save Failed", f"Could not save DOCX:\n{e}")

    def _print_document(self):
        if not self._current_doc or not self._current_doc.get("records"):
            QMessageBox.information(self, "No Document", "Load a document first.")
            return
        try:
            from PyQt5.QtPrintSupport import QPrinter, QPrintDialog
            from PyQt5.QtGui import QTextDocument

            printer = QPrinter(QPrinter.HighResolution)
            dialog  = QPrintDialog(printer, self)
            if dialog.exec_() != QPrintDialog.Accepted:
                return

            html = self._build_html_for_print()
            doc  = QTextDocument()
            doc.setHtml(html)
            doc.print_(printer)
        except Exception as e:
            QMessageBox.critical(self, "Print Failed", f"Could not print:\n{e}")

    def _build_html_for_print(self) -> str:
        """Build a clean HTML table from current records for PDF/print output."""
        records = self._current_doc.get("records", [])
        rows_html = ""
        for r in records:
            chem   = self._fmt_chem(r.get("chemical_use") or r.get("chemicals_use") or [])
            actual = self._fmt_chem(
                r.get("actual_chemicals_used") or r.get("actual_chemical_used") or [],
                key="actual_chemicals_name")
            cols = [
                (r.get("category") or "").capitalize() or "—",
                f"{r.get('month','')}/{r.get('date','')}/{r.get('year','')}",
                r.get("client_name", "") or "—",
                r.get("start_time", "") or "—",
                chem, actual,
                self._extract_remarks(r),
            ]
            cells = "".join(f"<td style='padding:6px 10px;border:1px solid #e0ede6'>{c}</td>"
                            for c in cols)
            rows_html += f"<tr>{cells}</tr>"

        headers = ["Category","Date","Client","Time","Chemicals Used","Actual Chemicals","Remarks"]
        hdr_cells = "".join(
            f"<th style='background:#2D6A4F;color:#fff;padding:8px 10px;"
            f"border:1px solid #1B4332;text-align:left'>{h}</th>"
            for h in headers)

        title = self._current_doc.get("title", "Inventory Report")
        return f"""
        <html><body style='font-family:Segoe UI,Arial;font-size:10pt;color:#1B4332'>
        <h2 style='color:#2D6A4F'>Raionn Pest Solutions</h2>
        <p style='color:#6B8F78'>Professional Pest Control and Inventory Management</p>
        <h3>{title}</h3>
        <table style='border-collapse:collapse;width:100%'>
            <tr>{hdr_cells}</tr>
            {rows_html}
        </table>
        </body></html>
        """

    # ── File list ─────────────────────────────────────────────────────────────

    def refresh_file_list(self):
        self.ui.filesList.clear()
        self._doc_map.clear()
        docs = self.api.get_documents()
        for doc in docs:
            self._add_file_item(doc)

    def _add_file_item(self, doc: dict):
        item_widget = QWidget()
        item_widget.setStyleSheet("background: transparent;")
        layout = QHBoxLayout(item_widget)
        layout.setContentsMargins(8, 4, 8, 4)
        layout.setSpacing(8)

        file_btn = QPushButton(f"📄  {doc['file_name']}")
        file_btn.setFlat(True)
        file_btn.setCursor(QCursor(Qt.PointingHandCursor))
        file_btn.setStyleSheet("""
            QPushButton { color: rgba(198,246,213,0.85); background: transparent;
                border: none; font: 10pt 'Segoe UI'; text-align: left; padding: 0; }
            QPushButton:hover { color: #FFFFFF; }
        """)
        file_btn.clicked.connect(lambda _, d=doc: self._load_document(d))
        layout.addWidget(file_btn, 1)

        del_btn = QPushButton("✕")
        del_btn.setFixedSize(20, 20)
        del_btn.setCursor(QCursor(Qt.PointingHandCursor))
        del_btn.setToolTip("Delete file")
        del_btn.setStyleSheet("""
            QPushButton { color: rgba(255,100,100,0.70); background: transparent;
                border: none; border-radius: 10px; font: 8pt 'Segoe UI'; padding: 0; }
            QPushButton:hover { color: #FF6B6B; background: rgba(255,80,80,0.18); }
            QPushButton:pressed { background: rgba(255,80,80,0.35); }
        """)
        del_btn.clicked.connect(lambda _, d=doc: self._delete_document(d))
        layout.addWidget(del_btn)

        list_item = QListWidgetItem()
        list_item.setSizeHint(QSize(item_widget.sizeHint().width(), 40))
        list_item.setData(Qt.UserRole, doc["id"])
        self.ui.filesList.addItem(list_item)
        self.ui.filesList.setItemWidget(list_item, item_widget)
        self._doc_map[doc["id"]] = doc

    def _delete_document(self, doc: dict):
        reply = QMessageBox.question(
            self, "Delete File",
            f"Delete '{doc['file_name']}'?\nThis action cannot be undone.",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            success, msg = self.api.delete_document(doc["id"])
            if success:
                self.refresh_file_list()
                if self._current_doc and self._current_doc.get("id") == doc["id"]:
                    self.ui.inventoryTable1.setRowCount(0)
                    self.ui.tocList.clear()
                    self.ui.docTitleLabel.setText("Inventory Report")
                    self._current_doc = None
                    self._raw_bytes   = None
            else:
                QMessageBox.critical(self, "Error", f"Failed to delete: {msg}")

    # ── Load document from API ────────────────────────────────────────────────

    def _on_file_selected(self, item: QListWidgetItem):
        doc_id = item.data(Qt.UserRole)
        doc    = self._doc_map.get(doc_id)
        if doc:
            self._load_document(doc)

    def _load_document(self, doc: dict):
        raw = self.api.download_document(doc["id"])
        if not raw:
            QMessageBox.warning(self, "Error", "Could not download document.")
            return
        self._raw_bytes = raw
        records = self._parse_docx_to_records(raw)
        self._current_doc = {"id": doc["id"], "title": doc["file_name"], "records": records}
        self._render_document(records, doc["file_name"])
        self._build_toc(records)

    # ── Render ────────────────────────────────────────────────────────────────

    def _render_document(self, records: list, title: str):
        t = self.ui.inventoryTable1
        t.setSortingEnabled(False)
        t.setRowCount(0)
        self.ui.docTitleLabel.setText(title)
        self.ui.metaLabel.setText(f"{len(records)} record(s) loaded")

        for r in records:
            row = t.rowCount()
            t.insertRow(row)
            values = [
                (r.get("category") or "").capitalize() or "—",
                f"{r.get('month','')}/{r.get('date','')}/{r.get('year','')}",
                r.get("client_name", "") or "—",
                r.get("start_time", "") or "—",
                self._fmt_chem(r.get("chemical_use") or r.get("chemicals_use") or []),
                self._fmt_chem(
                    r.get("actual_chemicals_used") or r.get("actual_chemical_used") or [],
                    key="actual_chemicals_name"),
                self._extract_remarks(r),
            ]
            for col, v in enumerate(values):
                item = QTableWidgetItem(str(v))
                item.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                t.setItem(row, col, item)

        t.setSortingEnabled(True)
        t.resizeRowsToContents()

    def _build_toc(self, records: list):
        self.ui.tocList.clear()
        self.ui.tocList.addItem("1. Inventory Table")
        for i, r in enumerate(records, start=1):
            date_str = f"{r.get('month','')}/{r.get('date','')}/{r.get('year','')}"
            client   = r.get("client_name", "—")
            category = (r.get("category", "") or "").capitalize()
            self.ui.tocList.addItem(
                QListWidgetItem(f"   {i}. [{category}] {client} — {date_str}"))
        self.ui.tocList.addItem("2. Statement")

    def _on_toc_selected(self, item: QListWidgetItem):
        text = item.text().strip()
        if text.startswith("1."):
            self.ui.pageScrollArea.verticalScrollBar().setValue(0)
        elif text.startswith("2."):
            self.ui.pageScrollArea.verticalScrollBar().setValue(
                self.ui.pageScrollArea.verticalScrollBar().maximum())

    # ── Helpers ───────────────────────────────────────────────────────────────

    @staticmethod
    def _extract_remarks(r: dict) -> str:
        top = (r.get("remarks") or "").strip()
        if top and top != "-":
            return top
        parts = []
        for c in (r.get("chemical_use") or r.get("chemicals_use") or []):
            val = (c.get("remarks") or "").strip()
            if val: parts.append(f"[C] {val}")
        for c in (r.get("actual_chemicals_used") or r.get("actual_chemical_used") or []):
            val = (c.get("remarks") or "").strip()
            if val: parts.append(f"[A.C.U] {val}")
        return "; ".join(parts) if parts else "—"

    @staticmethod
    def _fmt_chem(chems: list, key: str = "chemical_name") -> str:
        if not chems: return "—"
        parts = []
        for c in chems:
            name = (c.get(key) or c.get("chemical_name")
                    or c.get("actual_chemicals_name") or c.get("name") or "?")
            parts.append(name)
        return "; ".join(parts)

    @staticmethod
    def _parse_docx_to_records(raw: bytes) -> list:
        from io import BytesIO
        from docx import Document
        try:
            if isinstance(raw, str):
                raw = bytes.fromhex(raw)
            doc = Document(BytesIO(raw))
            if not doc.tables:
                return []
            tbl  = doc.tables[0]
            recs = []
            for row in tbl.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if not any(cells):
                    continue
                if len(cells) >= 7:
                    category, date_str, client, time_str, chem, actual, remarks = cells[:7]
                elif len(cells) == 6:
                    category = ""
                    date_str, client, time_str, chem, actual, remarks = cells[:6]
                else:
                    continue
                parts = date_str.split("/")
                month = parts[0] if len(parts) > 0 else ""
                day   = parts[1] if len(parts) > 1 else ""
                year  = parts[2] if len(parts) > 2 else ""
                recs.append({
                    "category":              category,
                    "month": month, "date": day, "year": year,
                    "client_name":           client,
                    "start_time":            time_str,
                    "chemical_use":          [{"chemical_name": chem}] if chem else [],
                    "actual_chemicals_used": [{"actual_chemicals_name": actual}] if actual else [],
                    "remarks":               remarks,
                })
            return recs
        except Exception:
            import traceback; traceback.print_exc()
            return []