from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ────────────────────────────────────────────────────────────────

def _fmt_time(record: dict) -> str:
    start = record.get("start_time", "") or ""
    end   = record.get("end_time",   "") or ""
    s_mer = record.get("start_meridiem", "") or ""
    e_mer = record.get("end_meridiem",   "") or ""

    start_full = f"{start} {s_mer}".strip()
    end_full   = f"{end} {e_mer}".strip()

    if start_full and end_full:
        return f"{start_full} - {end_full}"
    return start_full or end_full or "-"


def _fmt_chems(chems: list, name_key: str = "chemical_name") -> str:
    """Single-line semicolon-separated chemical names + quantities."""
    if not chems:
        return "-"
    parts = []
    for c in chems:
        name = (
            c.get(name_key)
            or c.get("chemical_name")
            or c.get("actual_chemicals_name")
            or c.get("name")
            or ""
        ).strip()
        qty = str(c.get("quantity") or c.get("qty") or "").strip()
        line = name
        if qty:
            line += f" ({qty})"
        if line.strip():
            parts.append(line)
    return "; ".join(parts) if parts else "-"


def _extract_remarks(record: dict) -> str:
    """
    Build a single remarks string with source labels:
      [C] <remark>       — from chemical_use entries
      [A.C.U] <remark>   — from actual_chemical_used entries

    If the record has a plain top-level "remarks" field (e.g. from a
    re-parsed docx), return it as-is so labels are never doubled.

    Returns "-" when nothing is found.
    """
    # ── Already-labelled string from a re-read docx ───────────────────────
    top = (record.get("remarks") or "").strip()
    if top and top != "-":
        return top

    parts = []

    # ── [C] — per-chemical remarks from chemical_use ──────────────────────
    chem_list = (
        record.get("chemical_use")
        or record.get("chemicals_use")
        or []
    )
    for c in chem_list:
        val = (c.get("remarks") or "").strip()
        if val:
            parts.append(f"[C] {val}")

    # ── [A.C.U] — per-chemical remarks from actual_chemical_used ──────────
    actual_list = (
        record.get("actual_chemical_used")
        or record.get("actual_chemicals_used")
        or []
    )
    for c in actual_list:
        val = (c.get("remarks") or "").strip()
        if val:
            parts.append(f"[A.C.U] {val}")

    return "; ".join(parts) if parts else "-"


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _write_cell(cell, text: str, font_size: int = 9,
                bold: bool = False, centered: bool = False,
                color: str = None):
    """Clear the cell and write exactly one clean run."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    p.clear()
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold      = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


# ── public API ─────────────────────────────────────────────────────────────

def generate_docx(records: list, title: str = "Inventory Report") -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x2E)

    doc.add_paragraph()

    HEADERS = [
        "Category", "Date", "Name of Client", "Time",
        "Chemical/s Used", "Actual Chemical/s Used", "Remarks",
    ]

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"

    for col_idx, header_text in enumerate(HEADERS):
        cell = table.rows[0].cells[col_idx]
        _set_cell_bg(cell, "1A562E")
        _write_cell(cell, header_text, font_size=9,
                    bold=True, centered=True, color="FFFFFF")

    for rec in records:
        category = (rec.get("category") or "").strip().capitalize() or "-"
        month    = str(rec.get("month", "") or "")
        day      = str(rec.get("date",  "") or "")
        year     = str(rec.get("year",  "") or "")
        date_str = "/".join(filter(None, [month, day, year])) or "-"
        client   = (rec.get("client_name") or "").strip() or "-"
        time_str = _fmt_time(rec)

        chem_list   = rec.get("chemical_use") or rec.get("chemicals_use") or []
        actual_list = rec.get("actual_chemical_used") or rec.get("actual_chemicals_used") or []

        chem_str   = _fmt_chems(chem_list,   name_key="chemical_name")
        actual_str = _fmt_chems(actual_list, name_key="chemical_name")

        # Remarks written with [C] / [A.C.U] labels into col 6
        remarks = _extract_remarks(rec)

        row = table.add_row()
        for col_idx, value in enumerate([
            category, date_str, client, time_str,
            chem_str, actual_str, remarks,
        ]):
            _write_cell(row.cells[col_idx], value, font_size=9)

    doc.add_paragraph()
    stmt = doc.add_heading("Statement", level=2)
    for run in stmt.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x2E)
    doc.add_paragraph(
        "This document is a system-generated inventory report. "
        "All records reflect data submitted through the inventory management system."
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()